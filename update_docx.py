from docx import Document

doc = Document('report.docx')

# Expand 'Branch Addressing' in Challenges
for p in doc.paragraphs:
    if p.text.startswith("Branch Addressing:"):
        p.text = p.text + " Because branch offsets specify the number of instructions to skip relative to the instruction immediately following the branch (PC+4), tracking the simulated program counter accurately across the entire file was highly complex. The 2-pass algorithm effectively solves this by pre-computing all absolute target addresses during Pass 1 without formatting the output, ensuring perfect alignment when string substitution and pseudo-instruction merging occurs in Pass 2."

# Add "6. Data Structures"
doc.add_heading('6. Data Structures', level=1)
doc.add_paragraph("To efficiently handle the disassembly process, we utilized native Python data structures. Specifically, a Dictionary is used for label mapping (mapping target memory addresses to string labels like 'loop_1' or 'func_1'), providing rapid lookups during disassembly. Additionally, Lists are used for binary instruction storage and accumulating the final assembly string output, allowing for ordered processing and easy post-processing modification.")

# Add "7. Array and Loop Reconstruction"
doc.add_heading('7. Array and Loop Reconstruction', level=1)
doc.add_paragraph("A sophisticated feature of our engine is its ability to interpret sequences of low-level instructions and infer high-level code structures. By pattern-matching sequences such as 'sll', 'addu', and 'lw', the engine dynamically reconstructs array access operations and annotates the output with the identified base and index registers. Furthermore, the engine structurally distinguishes between backward loops and forward procedure calls, assigning descriptive labels ('loop_X' and 'func_X' respectively) to fully reconstruct loops, conditional statements, and functions.")

doc.save('report.docx')
